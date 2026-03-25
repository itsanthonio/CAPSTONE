from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(14)

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.size = Pt(12)

def para(text):
    doc.add_paragraph(text)

def auth_line(text):
    p = doc.add_paragraph()
    r1 = p.add_run('Authentication: ')
    r1.bold = True
    r2 = p.add_run(text)
    r2.italic = True

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()

# ── Title ─────────────────────────────────────────────────────────────────
t = doc.add_heading('Appendix B: API Documentation — Changes Summary', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

para(
    'This document contains only the sections of the API documentation that changed '
    'under the role redesign. Replace the corresponding sections in the original document.'
)

# ── Updated intro ─────────────────────────────────────────────────────────
heading1('Updated Introduction (replace second sentence)')
para(
    'Three roles govern access: System Administrator — global access; manages users, '
    'organisations, model configuration, GEE quotas, and global reference data. '
    'Agency Administrator — organisation-scoped; submits scans, manages alerts, assigns '
    'inspectors, and views results for their organisation only. '
    'Inspector — organisation-scoped; views their own assignments and resolves alerts '
    'assigned to them.'
)

# ── Alert write endpoints ─────────────────────────────────────────────────
heading1('6. Alerts (changed endpoints only)')

heading2('POST   /api/alerts/')
para('Manually create an alert at a given location. Automatically creates an associated DetectedSite record.')
auth_line('Agency Admin or System Admin (IsAdminRole)')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['title', 'string', 'Required', 'Short descriptive title for the alert.'],
        ['latitude', 'float', 'Required', 'WGS 84 latitude of the detection point.'],
        ['longitude', 'float', 'Required', 'WGS 84 longitude of the detection point.'],
        ['severity', 'string', 'Required', 'Values: critical, high, medium, low.'],
        ['alert_type', 'string', 'Required', 'Values: mining_activity, environmental_risk.'],
        ['description', 'string', 'Optional', 'Longer free-text description.'],
        ['confidence_score', 'float', 'Optional', 'Model confidence 0.0–1.0.'],
        ['area_hectares', 'float', 'Optional', 'Estimated affected area. Default: 0.01.'],
        ['legal_status', 'string', 'Optional', 'Values: unknown, legal, illegal. Default: unknown.'],
        ['region_id', 'UUID', 'Optional', 'Associate with an existing Region record.'],
        ['assigned_to_id', 'string', 'Optional', 'Username of the inspector to assign immediately.'],
        ['status', 'string', 'Optional', 'Initial status. Values: open, dispatched. Default: open.'],
    ]
)

heading2('POST   /api/alerts/{id}/acknowledge/')
para('Mark an alert as acknowledged. Records the timestamp and acting user.')
auth_line('Agency Admin or System Admin')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the alert.']])

heading2('POST   /api/alerts/{id}/dismiss/')
para('Dismiss an alert, removing it from the active queue.')
auth_line('Agency Admin or System Admin')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the alert.']])

heading2('POST   /api/alerts/{id}/dispatch/')
para('Mark an alert as dispatched for field inspection.')
auth_line('Agency Admin or System Admin')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the alert.']])

heading2('POST   /api/alerts/{id}/assign_inspector/')
para('Assign a field inspector to an alert. Automatically transitions the alert status to dispatched.')
auth_line('Agency Admin or System Admin')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [
        ['id', 'UUID (path)', 'Required', 'UUID of the alert.'],
        ['inspector_id', 'UUID', 'Required', 'UserProfile.id of the inspector to assign.'],
    ])

heading2('PUT   /api/alerts/{id}/update/')
para('Replace all editable fields on an alert record. Use PATCH for partial updates.')
auth_line('Agency Admin or System Admin')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [
        ['id', 'UUID (path)', 'Required', '—'],
        ['title', 'string', 'Optional', '—'],
        ['description', 'string', 'Optional', '—'],
        ['severity', 'string', 'Optional', 'Values: critical, high, medium, low.'],
        ['alert_type', 'string', 'Optional', 'Values: mining_activity, environmental_risk.'],
        ['status', 'string', 'Optional', 'Values: open, acknowledged, dispatched, resolved, dismissed.'],
        ['assigned_to_id', 'string', 'Optional', 'Username, or null to unassign.'],
        ['latitude', 'float', 'Optional', 'Updated location latitude.'],
        ['longitude', 'float', 'Optional', 'Updated location longitude.'],
        ['confidence_score', 'float', 'Optional', '0.0–1.0.'],
        ['area_hectares', 'float', 'Optional', '—'],
        ['resolution_notes', 'string', 'Optional', '—'],
    ])

heading2('PATCH   /api/alerts/{id}/update/')
para('Partially update one or more fields on an alert record.')
auth_line('Agency Admin or System Admin')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [
        ['id', 'UUID (path)', 'Required', '—'],
        ['title', 'string', 'Optional', '—'],
        ['description', 'string', 'Optional', '—'],
        ['severity', 'string', 'Optional', 'Values: critical, high, medium, low.'],
        ['alert_type', 'string', 'Optional', 'Values: mining_activity, environmental_risk.'],
        ['status', 'string', 'Optional', 'Values: open, acknowledged, dispatched, resolved, dismissed.'],
        ['assigned_to_id', 'string', 'Optional', 'Username, or null to unassign.'],
        ['latitude', 'float', 'Optional', 'Updated location latitude.'],
        ['longitude', 'float', 'Optional', 'Updated location longitude.'],
        ['confidence_score', 'float', 'Optional', '0.0–1.0.'],
        ['area_hectares', 'float', 'Optional', '—'],
        ['resolution_notes', 'string', 'Optional', '—'],
    ])

heading2('DELETE   /api/alerts/{id}/delete/')
para('Permanently delete an alert record. This action is irreversible.')
auth_line('System Admin only')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the alert.']])

heading2('POST   /api/alerts/bulk_action/')
para('Apply a single status transition to multiple alerts in one request.')
auth_line('Agency Admin or System Admin')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [
        ['ids', 'UUID[] (body)', 'Required', 'Array of alert UUIDs to update.'],
        ['action', 'string (body)', 'Required', 'Transition to apply. Values: acknowledged, dismissed, dispatched.'],
    ])

heading2('POST   /api/alerts/bulk-assign-inspector/')
para('Assign the same field inspector to multiple alerts simultaneously. All affected alerts are transitioned to dispatched.')
auth_line('Agency Admin or System Admin')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [
        ['alert_ids', 'UUID[] (body)', 'Required', 'Array of alert UUIDs.'],
        ['inspector_username', 'string (body)', 'Required', 'Username of the inspector to assign.'],
    ])

# ── Data Uploads ──────────────────────────────────────────────────────────
heading1('7. Data Uploads (all three endpoints changed)')

heading2('POST   /uploads/upload/concessions/')
para(
    'Import legal mining concession boundaries from a GeoJSON file. Each feature must include '
    'the properties: license_number, concession_name, holder_name, and license_type. After '
    'import the system re-evaluates the legal status of all existing DetectedSite records '
    'against the new concession data.'
)
auth_line('System Admin only')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [['file', 'file (multipart/form-data)', 'Required', 'GeoJSON FeatureCollection (.geojson or .json).']])

heading2('POST   /uploads/upload/water-bodies/')
para('Import water body boundary polygons from a GeoJSON file. Used for environmental context layers on the map and for legal status evaluation of detected sites.')
auth_line('System Admin only')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [['file', 'file (multipart/form-data)', 'Required', 'GeoJSON FeatureCollection (.geojson or .json).']])

heading2('POST   /uploads/upload/protected-forests/')
para('Import protected forest boundary polygons from a GeoJSON file. Used for environmental context layers on the map and for legal status evaluation of detected sites.')
auth_line('System Admin only')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [['file', 'file (multipart/form-data)', 'Required', 'GeoJSON FeatureCollection (.geojson or .json).']])

# ── Scanning ──────────────────────────────────────────────────────────────
heading1('8. Scanning System (changed endpoints only)')

heading2('POST   /scanning/api/toggle/')
para('Pause or resume the automated tile scanner. When paused, the Celery beat task continues running but skips tile selection and job submission.')
auth_line('System Admin only')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [['action', 'string (body)', 'Required', 'Values: pause, resume.']])

heading2('POST   /scanning/api/force-scan/')
para('Immediately queue a scan job for a specific tile, bypassing the normal priority schedule. The tile is promoted to the front of the priority queue.')
auth_line('Agency Admin or System Admin')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [
        ['lat', 'float (body)', 'Required', 'WGS 84 latitude of the tile centre.'],
        ['lng', 'float (body)', 'Required', 'WGS 84 longitude of the tile centre.'],
    ])

# ── New sections ──────────────────────────────────────────────────────────
heading1('9. User Management (new section)')

para('These endpoints are accessible to System Admins only and are used to manage user accounts and organisation assignments across the platform.')

heading2('GET   /api/users/')
para('List all user accounts. Results are paginated.')
auth_line('System Admin only')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [
        ['page', 'integer (query)', 'Optional', 'Page number. Default: 1.'],
        ['role', 'string (query)', 'Optional', 'Filter by role. Values: system_admin, agency_admin, inspector.'],
        ['organisation_id', 'UUID (query)', 'Optional', 'Filter by organisation.'],
    ])

heading2('POST   /api/users/')
para('Create a new user account.')
auth_line('System Admin only')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [
        ['username', 'string', 'Required', 'Unique username.'],
        ['email', 'string', 'Required', 'User email address.'],
        ['password', 'string', 'Required', 'Initial password.'],
        ['role', 'string', 'Required', 'Values: system_admin, agency_admin, inspector.'],
        ['organisation_id', 'UUID', 'Optional', 'Required for agency_admin and inspector roles.'],
        ['first_name', 'string', 'Optional', '—'],
        ['last_name', 'string', 'Optional', '—'],
    ])

heading2('GET   /api/users/{id}/')
para('Retrieve full details for a single user account.')
auth_line('System Admin only')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the user profile.']])

heading2('PATCH   /api/users/{id}/')
para('Partially update a user account, including role or organisation assignment.')
auth_line('System Admin only')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [
        ['id', 'UUID (path)', 'Required', '—'],
        ['role', 'string', 'Optional', 'Values: system_admin, agency_admin, inspector.'],
        ['organisation_id', 'UUID', 'Optional', 'Reassign the user to a different organisation.'],
        ['is_active', 'boolean', 'Optional', 'Set to false to deactivate the account.'],
        ['first_name', 'string', 'Optional', '—'],
        ['last_name', 'string', 'Optional', '—'],
        ['email', 'string', 'Optional', '—'],
    ])

heading1('10. Organisation Management (new section)')

para('These endpoints are accessible to System Admins only and are used to manage the organisations (agencies) registered on the platform.')

heading2('GET   /api/organisations/')
para('List all registered organisations.')
auth_line('System Admin only')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [['page', 'integer (query)', 'Optional', 'Page number. Default: 1.']])

heading2('POST   /api/organisations/')
para('Create a new organisation.')
auth_line('System Admin only')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [
        ['name', 'string', 'Required', 'Display name of the organisation.'],
        ['district_scope', 'string[]', 'Optional', 'List of ADM2_NAME district names this organisation is scoped to.'],
    ])

heading2('PATCH   /api/organisations/{id}/')
para('Update the name or district scope of an organisation.')
auth_line('System Admin only')
add_table(['Field', 'Type', 'Required', 'Notes'],
    [
        ['id', 'UUID (path)', 'Required', 'UUID of the organisation.'],
        ['name', 'string', 'Optional', '—'],
        ['district_scope', 'string[]', 'Optional', 'Replaces the existing district scope list.'],
    ])

out = r'C:\Users\mcnob\Documents\Ashesi A\Cappy Cap\api_documentation_changes.docx'
doc.save(out)
print(f'Saved to {out}')
