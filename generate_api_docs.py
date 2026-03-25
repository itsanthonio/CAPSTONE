from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(14)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.size = Pt(12)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.size = Pt(11)
    return p

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def auth_line(text):
    p = doc.add_paragraph()
    run_label = p.add_run('Authentication: ')
    run_label.bold = True
    run_val = p.add_run(text)
    run_val.italic = True
    return p

def scope_note(text):
    p = doc.add_paragraph()
    run = p.add_run('Scoping note: ')
    run.bold = True
    run2 = p.add_run(text)
    run2.italic = True
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()  # spacing after table

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_heading('Appendix B: API Documentation', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

para(
    'This document describes all REST API endpoints exposed by the SankofaWatch platform. '
    'Endpoints are grouped by resource. All authenticated endpoints require a valid session '
    'cookie or token.\n\n'
    'Three roles govern access:\n'
    '  • System Administrator — global access; manages users, organisations, model configuration, '
    'GEE quotas, and global reference data.\n'
    '  • Agency Administrator — organisation-scoped; submits scans, manages alerts, assigns '
    'inspectors, and views results for their organisation only.\n'
    '  • Inspector — organisation-scoped; views their own assignments and resolves alerts '
    'assigned to them.\n\n'
    'For each endpoint, the parameter table describes the fields that must be included in the '
    'request body or query string. Required fields must always be provided, while optional '
    'fields have default values where specified.'
)

# ── 1. Jobs ──────────────────────────────────────────────────────────────────
heading1('1. Jobs')

heading2('POST   /api/jobs/')
para(
    'Submit a new detection job for an area of interest. Processing is asynchronous; the '
    'response returns immediately with the queued job record. The Celery task ID and Job ID '
    'are included in the response headers X-Task-ID and X-Job-ID, respectively.'
)
auth_line('IsAuthenticated — rate-limited per user (JobCreateThrottle)')
scope_note('Agency Admins may only submit jobs within their organisation\'s assigned districts.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['aoi_geometry', 'GeoJSON object', 'Required', 'Polygon or MultiPolygon. Area must be between 100 ha and 1,000 ha. Maximum 1,000 vertices.'],
        ['start_date', 'string', 'Required', 'ISO date (YYYY-MM-DD). Must be earlier than end_date.'],
        ['end_date', 'string', 'Required', 'ISO date (YYYY-MM-DD).'],
        ['model_version', 'string', 'Optional', 'ML model version to use. Default: "v1.0".'],
    ]
)

heading2('GET   /api/jobs/')
para('List all detection jobs. Results are paginated.')
auth_line('IsAuthenticated')
scope_note('Agency Admins receive only jobs submitted within their organisation. System Admins receive all jobs.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['page', 'integer', 'Optional', 'Page number. Default: 1.'],
        ['page_size', 'integer', 'Optional', 'Items per page. Default: system setting.'],
    ]
)

heading2('GET   /api/jobs/{id}/')
para('Retrieve the full record for a single detection job.')
auth_line('IsAuthenticated')
scope_note('Agency Admins may only retrieve jobs belonging to their organisation.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the job.']]
)

heading2('GET   /api/jobs/{id}/status/')
para(
    'Poll the processing status and pipeline progress of a job. Use this endpoint to track '
    'a job after submission, rather than repeatedly fetching the full record.'
)
auth_line('IsAuthenticated')
scope_note('Agency Admins may only poll jobs belonging to their organisation.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the job.']]
)

# ── 2. Results ───────────────────────────────────────────────────────────────
heading1('2. Results')

heading2('GET   /api/results/')
para('List detection results. Optionally filter by job.')
auth_line('IsAuthenticated')
scope_note('Agency Admins receive only results from jobs within their organisation.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['job_id', 'UUID (query)', 'Optional', 'Filter results belonging to this job.']]
)

heading2('GET   /api/results/{id}/')
para('Retrieve a single result record including the full GeoJSON detection output.')
auth_line('IsAuthenticated')
scope_note('Agency Admins may only retrieve results from their organisation\'s jobs.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the result.']]
)

heading2('GET   /api/results/{job_id}/by_job/')
para('Retrieve all results associated with a specific job in a single response. Only available once the job status is \'completed\'.')
auth_line('IsAuthenticated')
scope_note('Agency Admins may only access results from their organisation\'s jobs.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['job_id', 'UUID (path)', 'Required', 'UUID of the parent job.']]
)

# ── 3. Detected Sites ────────────────────────────────────────────────────────
heading1('3. Detected Sites')

heading2('GET   /api/sites/')
para('List all detected mining sites as a GeoJSON FeatureCollection. Supports pagination for large datasets. The default limit is 2,000 features.')
auth_line('IsAuthenticated')
scope_note('Agency Admins receive only sites detected within their organisation\'s district scope. System Admins receive all sites.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['page', 'integer (query)', 'Optional', 'Page number. Default: 1.'],
        ['per_page', 'integer (query)', 'Optional', 'Features per page. Default: 200, maximum implied limit: 2,000.'],
    ]
)

heading2('GET   /api/sites/{id}/')
para('Retrieve full details for a single detected site, including concession overlap information and patch image URLs.')
auth_line('IsAuthenticated')
scope_note('Agency Admins may only retrieve sites within their organisation\'s scope.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the detected site.']]
)

heading2('GET   /api/sites/{id}/timelapse/')
para('Retrieve historical satellite thumbnail frames for a detected site. Frames cover up to five dry seasons (November–March) prior to the detection date.')
auth_line('IsAuthenticated')
scope_note('Agency Admins may only access timelapse data for sites within their organisation\'s scope.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the detected site.']]
)

# ── 4. Concessions ───────────────────────────────────────────────────────────
heading1('4. Concessions')

heading2('GET   /api/concessions/')
para('List all legal mining concessions as a GeoJSON FeatureCollection. Each feature includes the license number, concession name, holder, and license type.')
auth_line('IsAuthenticated')
para('No parameters.')

# ── 5. Regions ───────────────────────────────────────────────────────────────
heading1('5. Regions')

heading2('GET   /api/regions/')
para('List protected region boundaries (water bodies and protected forests) as a GeoJSON FeatureCollection. Optionally filter by region type.')
auth_line('IsAuthenticated')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['type', 'string (query)', 'Optional', 'Filter by region type. Accepted values: water_body, protected_forest.']]
)

# ── 6. Alerts ────────────────────────────────────────────────────────────────
heading1('6. Alerts')

heading2('GET   /api/alerts/')
para('List alerts with filtering and ordering options. Results are paginated.')
auth_line('IsAuthenticated')
scope_note('Agency Admins receive only alerts within their organisation\'s district scope. System Admins receive all alerts.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['status', 'string (query)', 'Optional', 'Filter by alert status. Values: open, acknowledged, dispatched, resolved, dismissed.'],
        ['severity', 'string (query)', 'Optional', 'Filter by severity. Values: critical, high, medium, low.'],
        ['alert_type', 'string (query)', 'Optional', 'Filter by type. Values: mining_activity, environmental_risk.'],
        ['source', 'string (query)', 'Optional', 'Filter by origin. Values: manual, automated.'],
        ['page', 'integer (query)', 'Optional', 'Page number. Default: 1.'],
        ['per_page', 'integer (query)', 'Optional', 'Items per page. Default: 20.'],
        ['ordering', 'string (query)', 'Optional', 'Sort field. Example: -created_at, -detected_site__confidence_score.'],
    ]
)

heading2('GET   /api/alerts/{id}/')
para('Retrieve full detail for a single alert including associated site data and field verification record if one exists.')
auth_line('IsAuthenticated')
scope_note('Agency Admins may only retrieve alerts within their organisation\'s scope.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the alert.']]
)

heading2('POST   /api/alerts/')
para('Manually create an alert at a given location. Automatically creates an associated DetectedSite record.')
auth_line('Agency Admin or System Admin (IsAdminRole)')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['title', 'string', 'Required', 'Short descriptive title for the alert.'],
        ['latitude', 'float', 'Required', 'WGS 84 latitude of the detection point.'],
        ['longitude', 'float', 'Required', 'WGS 84 longitude of the detection point.'],
        ['severity', 'string', 'Required', 'Values: critical, high, medium, low.'],
        ['alert_type', 'string', 'Required', 'Values: mining_activity, environmental_risk.'],
        ['description', 'string', 'Optional', 'Longer free-text description.'],
        ['confidence_score', 'float', 'Optional', 'Model confidence 0.0–1.0.'],
        ['area_hectares', 'float', 'Optional', 'Estimated affected area. Default: 0.01.'],
        ['legal_status', 'string', 'Optional', 'Values: unknown, legal, illegal. Default: unknown.'],
        ['region_id', 'UUID', 'Optional', 'Associate with an existing Region record.'],
        ['assigned_to_id', 'string', 'Optional', 'Username of the inspector to assign immediately.'],
        ['status', 'string', 'Optional', 'Initial status. Values: open, dispatched. Default: open.'],
    ]
)

heading2('GET   /api/alerts/summary/')
para('Return aggregate counts of alerts grouped by status and severity. Useful for dashboard KPI widgets.')
auth_line('IsAuthenticated')
scope_note('Agency Admins receive counts scoped to their organisation only.')
para('No parameters.')

heading2('POST   /api/alerts/{id}/acknowledge/')
para('Mark an alert as acknowledged. Records the timestamp and acting user.')
auth_line('Agency Admin or System Admin')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the alert.']]
)

heading2('POST   /api/alerts/{id}/dismiss/')
para('Dismiss an alert, removing it from the active queue.')
auth_line('Agency Admin or System Admin')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the alert.']]
)

heading2('POST   /api/alerts/{id}/dispatch/')
para('Mark an alert as dispatched for field inspection.')
auth_line('Agency Admin or System Admin')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the alert.']]
)

heading2('POST   /api/alerts/{id}/resolve/')
para('Mark an alert as resolved. Accessible to admins and the inspector currently assigned to the alert.')
auth_line('IsAuthenticated (Agency Admin, System Admin, or assigned Inspector)')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['id', 'UUID (path)', 'Required', 'UUID of the alert.'],
        ['resolution_notes', 'string', 'Optional', 'Free-text summary of the resolution outcome.'],
    ]
)

heading2('POST   /api/alerts/{id}/assign_inspector/')
para('Assign a field inspector to an alert. Automatically transitions the alert status to dispatched.')
auth_line('Agency Admin or System Admin')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['id', 'UUID (path)', 'Required', 'UUID of the alert.'],
        ['inspector_id', 'UUID', 'Required', 'UserProfile.id of the inspector to assign.'],
    ]
)

heading2('PUT   /api/alerts/{id}/update/')
para('Replace all editable fields on an alert record. Use PATCH for partial updates.')
auth_line('Agency Admin or System Admin')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['id', 'UUID (path)', 'Required', '—'],
        ['title', 'string', 'Optional', '—'],
        ['description', 'string', 'Optional', '—'],
        ['severity', 'string', 'Optional', 'Values: critical, high, medium, low.'],
        ['alert_type', 'string', 'Optional', 'Values: mining_activity, environmental_risk.'],
        ['status', 'string', 'Optional', 'Values: open, acknowledged, dispatched, resolved, dismissed.'],
        ['assigned_to_id', 'string', 'Optional', 'Username, or null to unassign.'],
        ['latitude', 'float', 'Optional', 'Updated location latitude.'],
        ['longitude', 'float', 'Optional', 'Updated location longitude.'],
        ['confidence_score', 'float', 'Optional', '0.0–1.0.'],
        ['area_hectares', 'float', 'Optional', '—'],
        ['resolution_notes', 'string', 'Optional', '—'],
    ]
)

heading2('PATCH   /api/alerts/{id}/update/')
para('Partially update one or more fields on an alert record.')
auth_line('Agency Admin or System Admin')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['id', 'UUID (path)', 'Required', '—'],
        ['title', 'string', 'Optional', '—'],
        ['description', 'string', 'Optional', '—'],
        ['severity', 'string', 'Optional', 'Values: critical, high, medium, low.'],
        ['alert_type', 'string', 'Optional', 'Values: mining_activity, environmental_risk.'],
        ['status', 'string', 'Optional', 'Values: open, acknowledged, dispatched, resolved, dismissed.'],
        ['assigned_to_id', 'string', 'Optional', 'Username, or null to unassign.'],
        ['latitude', 'float', 'Optional', 'Updated location latitude.'],
        ['longitude', 'float', 'Optional', 'Updated location longitude.'],
        ['confidence_score', 'float', 'Optional', '0.0–1.0.'],
        ['area_hectares', 'float', 'Optional', '—'],
        ['resolution_notes', 'string', 'Optional', '—'],
    ]
)

heading2('DELETE   /api/alerts/{id}/delete/')
para('Permanently delete an alert record. This action is irreversible.')
auth_line('System Admin only')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the alert.']]
)

heading2('POST   /api/alerts/bulk_action/')
para('Apply a single status transition to multiple alerts in one request.')
auth_line('Agency Admin or System Admin')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['ids', 'UUID[] (body)', 'Required', 'Array of alert UUIDs to update.'],
        ['action', 'string (body)', 'Required', 'Transition to apply. Values: acknowledged, dismissed, dispatched.'],
    ]
)

heading2('POST   /api/alerts/bulk-assign-inspector/')
para('Assign the same field inspector to multiple alerts simultaneously. All affected alerts are transitioned to dispatched.')
auth_line('Agency Admin or System Admin')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['alert_ids', 'UUID[] (body)', 'Required', 'Array of alert UUIDs.'],
        ['inspector_username', 'string (body)', 'Required', 'Username of the inspector to assign.'],
    ]
)

heading2('GET   /api/my-assignments/')
para('Return the pending inspection assignments for the currently authenticated user. Intended for inspector dashboards.')
auth_line('IsAuthenticated (login required)')
para('No parameters.')

# ── 7. Data Uploads ──────────────────────────────────────────────────────────
heading1('7. Data Uploads')

heading2('POST   /uploads/upload/concessions/')
para(
    'Import legal mining concession boundaries from a GeoJSON file. Each feature must include '
    'the properties: license_number, concession_name, holder_name, and license_type. After '
    'import the system re-evaluates the legal status of all existing DetectedSite records '
    'against the new concession data.'
)
auth_line('System Admin only')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['file', 'file (multipart/form-data)', 'Required', 'GeoJSON FeatureCollection (.geojson or .json).']]
)

heading2('POST   /uploads/upload/water-bodies/')
para('Import water body boundary polygons from a GeoJSON file. Used for environmental context layers on the map and for legal status evaluation of detected sites.')
auth_line('System Admin only')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['file', 'file (multipart/form-data)', 'Required', 'GeoJSON FeatureCollection (.geojson or .json).']]
)

heading2('POST   /uploads/upload/protected-forests/')
para('Import protected forest boundary polygons from a GeoJSON file. Used for environmental context layers on the map and for legal status evaluation of detected sites.')
auth_line('System Admin only')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['file', 'file (multipart/form-data)', 'Required', 'GeoJSON FeatureCollection (.geojson or .json).']]
)

# ── 8. Scanning System ───────────────────────────────────────────────────────
heading1('8. Scanning System')

heading2('GET   /scanning/api/status/')
para(
    'Return a full snapshot of the automated scanning system state, including whether the '
    'system is active, today\'s scan and detection counts, currently running jobs, the next '
    'tiles queued, and historical detection aggregations by day and region.'
)
auth_line('IsAuthenticated (login required)')
scope_note('Agency Admins receive counts and tile data scoped to their organisation\'s districts. System Admins receive the global snapshot.')
para('No parameters.')

heading2('POST   /scanning/api/toggle/')
para('Pause or resume the automated tile scanner. When paused, the Celery beat task continues running but skips tile selection and job submission.')
auth_line('System Admin only')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['action', 'string (body)', 'Required', 'Values: pause, resume.']]
)

heading2('GET   /scanning/api/recent-tiles/')
para('Retrieve recently scanned tile geometries as a GeoJSON FeatureCollection. Useful for rendering the scanning grid animation on the auto scan map.')
auth_line('IsAuthenticated (login required)')
scope_note('Agency Admins receive only tiles within their organisation\'s district scope.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['limit', 'integer (query)', 'Optional', 'Maximum number of tiles to return. Default: 100, maximum: 500.'],
        ['date', 'string (query)', 'Optional', 'Filter to tiles scanned on this date (YYYY-MM-DD). Default: today.'],
    ]
)

heading2('GET   /scanning/api/detections/')
para('Retrieve automated detection points as a GeoJSON FeatureCollection. Returns point geometries with confidence, area, legal status, and region properties.')
auth_line('IsAuthenticated (login required)')
scope_note('Agency Admins receive only detections within their organisation\'s district scope.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['limit', 'integer (query)', 'Optional', 'Maximum number of features to return. Default: 200, maximum: 1,000.'],
        ['date', 'string (query)', 'Optional', 'Filter to detections created on this date (YYYY-MM-DD). Default: today.'],
    ]
)

heading2('GET   /scanning/api/tile-detail/')
para('Retrieve the scan history and associated detections for the tile at the given coordinates. Used to populate the tile info panel when the user clicks on the scanning grid.')
auth_line('IsAuthenticated (login required)')
scope_note('Agency Admins may only query tiles within their organisation\'s district scope.')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['lat', 'float (query)', 'Required', 'WGS 84 latitude of the tile centre.'],
        ['lng', 'float (query)', 'Required', 'WGS 84 longitude of the tile centre.'],
    ]
)

heading2('POST   /scanning/api/force-scan/')
para('Immediately queue a scan job for a specific tile, bypassing the normal priority schedule. The tile is promoted to the front of the priority queue.')
auth_line('Agency Admin or System Admin')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['lat', 'float (body)', 'Required', 'WGS 84 latitude of the tile centre.'],
        ['lng', 'float (body)', 'Required', 'WGS 84 longitude of the tile centre.'],
    ]
)

# ── 9. User Management (System Admin) ────────────────────────────────────────
heading1('9. User Management')

para('These endpoints are accessible to System Admins only and are used to manage user accounts and organisation assignments across the platform.')

heading2('GET   /api/users/')
para('List all user accounts. Results are paginated.')
auth_line('System Admin only')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['page', 'integer (query)', 'Optional', 'Page number. Default: 1.'],
        ['role', 'string (query)', 'Optional', 'Filter by role. Values: system_admin, agency_admin, inspector.'],
        ['organisation_id', 'UUID (query)', 'Optional', 'Filter by organisation.'],
    ]
)

heading2('POST   /api/users/')
para('Create a new user account.')
auth_line('System Admin only')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['username', 'string', 'Required', 'Unique username.'],
        ['email', 'string', 'Required', 'User email address.'],
        ['password', 'string', 'Required', 'Initial password.'],
        ['role', 'string', 'Required', 'Values: system_admin, agency_admin, inspector.'],
        ['organisation_id', 'UUID', 'Optional', 'Required for agency_admin and inspector roles.'],
        ['first_name', 'string', 'Optional', '—'],
        ['last_name', 'string', 'Optional', '—'],
    ]
)

heading2('GET   /api/users/{id}/')
para('Retrieve full details for a single user account.')
auth_line('System Admin only')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['id', 'UUID (path)', 'Required', 'UUID of the user profile.']]
)

heading2('PATCH   /api/users/{id}/')
para('Partially update a user account, including role or organisation assignment.')
auth_line('System Admin only')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['id', 'UUID (path)', 'Required', '—'],
        ['role', 'string', 'Optional', 'Values: system_admin, agency_admin, inspector.'],
        ['organisation_id', 'UUID', 'Optional', 'Reassign the user to a different organisation.'],
        ['is_active', 'boolean', 'Optional', 'Set to false to deactivate the account.'],
        ['first_name', 'string', 'Optional', '—'],
        ['last_name', 'string', 'Optional', '—'],
        ['email', 'string', 'Optional', '—'],
    ]
)

# ── 10. Organisation Management (System Admin) ───────────────────────────────
heading1('10. Organisation Management')

para('These endpoints are accessible to System Admins only and are used to manage the organisations (agencies) registered on the platform.')

heading2('GET   /api/organisations/')
para('List all registered organisations.')
auth_line('System Admin only')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [['page', 'integer (query)', 'Optional', 'Page number. Default: 1.']]
)

heading2('POST   /api/organisations/')
para('Create a new organisation.')
auth_line('System Admin only')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['name', 'string', 'Required', 'Display name of the organisation.'],
        ['district_scope', 'string[]', 'Optional', 'List of ADM2_NAME district names this organisation is scoped to.'],
    ]
)

heading2('PATCH   /api/organisations/{id}/')
para('Update the name or district scope of an organisation.')
auth_line('System Admin only')
add_table(
    ['Field', 'Type', 'Required', 'Notes'],
    [
        ['id', 'UUID (path)', 'Required', 'UUID of the organisation.'],
        ['name', 'string', 'Optional', '—'],
        ['district_scope', 'string[]', 'Optional', 'Replaces the existing district scope list.'],
    ]
)

# ── Save ─────────────────────────────────────────────────────────────────────
out = r'C:\Users\mcnob\Documents\Ashesi A\Cappy Cap\api_documentation.docx'
doc.save(out)
print(f'Saved to {out}')
